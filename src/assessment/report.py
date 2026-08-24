"""WP-1.7 — report rendering.

Three artifacts from one bundle:

* **JSON** — canonical (sorted keys, no insignificant whitespace, UTF-8). This
  is the ledger artifact, and ``report_hash`` is ``sha256`` of exactly these
  bytes.
* **Markdown** — the human-readable report.
* **DOCX** — rendered with ``python-docx`` (chosen over a Node ``docx``
  toolchain so the assessment CLI has no JavaScript runtime dependency). If
  ``python-docx`` is not installed, ``render()`` returns no ``docx`` path and
  states the reason; it never emits an empty file to make the output look
  complete.

TEMPLATE DISCIPLINE (linted by ``tests/assessment/test_report_lint.py``)
-----------------------------------------------------------------------

Every template string lives in :data:`TEMPLATES`, and no template contains a
numeric literal other than a section number in a heading. Numbers reach the
page only by substitution from measured values.

Every table of numbers in the rendered body either carries ``Grade`` and
``Provenance`` columns, or is immediately preceded by a ``**Grade:**`` and a
``**Provenance:**`` line covering every number in it (R9). The lint asserts
this over the *rendered* document, so a table added later without provenance
fails the build.

The report never states an engagement fee, a timeline, a delivery date, or a
behavioural-equivalence claim. It reports coverage and complexity; quoting is a
human step (R11). ``tests/assessment/test_report_lint.py`` asserts the absence
of that vocabulary too.
"""

from __future__ import annotations

import hashlib
import platform
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from . import complexity as complexity_mod
from . import coverage as coverage_mod
from . import loc as loc_mod
from .models import (
    AssessmentBundle,
    Measured,
    ProgramAssessment,
    canonical_json,
    measured_to_dict,
)
from .risk import programs_by_tier, rule_table, tier_counts
from .supported import (
    boundary_only_tokens,
    registry_provenance,
    supported_data_features,
    supported_verbs,
)

SCHEMA_VERSION = "relian-assessment-1"

TEMPLATES: Dict[str, str] = {
    "title": "# Legacy Code Assessment — {root}\n",
    "subtitle": (
        "Schema `{schema}` · manifest `{manifest_hash}`\n\n"
        "Every number in this report is a measurement with a stated origin and a "
        "Trutina grade, or it is absent. Nothing here is a default, an estimate, "
        "or a target reported as a result.\n"
    ),
    "h1": "\n## {n}. {title}\n\n",
    "h2": "\n### {title}\n\n",
    "grade_line": "**Grade:** {grade} · **Provenance:** {provenance}\n\n",
    "para": "{text}\n\n",
    "bullet": "- {text}\n",
    "table_row": "| {cells} |\n",
    "code_block": "```\n{text}\n```\n\n",
    "no_data": "_Not measured — {reason}._\n\n",

    # Section 0. Plain language, no jargon, and no number that is not restated
    # from a measurement below.
    "plain_intro": (
        "This section says what the rest of the report found, in ordinary words. "
        "It adds nothing: every figure here is repeated from a measured section "
        "below, and the section it came from is named so you can check it.\n"
    ),
    "plain_scope": (
        "**What was examined.** {programs} COBOL program(s) in `{root}`, out of "
        "{files} file(s) found. Nothing was changed and nothing left this "
        "machine — the assessment only reads. The exact list of files, with a "
        "checksum for each, is under **Manifest**.\n"
    ),
    "plain_scope_empty": (
        "**What was examined.** No COBOL programs were found in `{root}`, so "
        "there is nothing to report below. **Manifest** lists the files that were "
        "found.\n"
    ),
    "plain_verdict_head": (
        "**Where the code stands.** Each program falls into one of four groups "
        "(**Risk tiers** gives the tier for each program, and appendix C the "
        "rules that decide it):\n"
    ),
    "plain_blockers_head": (
        "**What stands in the way.** The COBOL features holding programs back, "
        "most common first; the **Unsupported-construct inventory** has the "
        "full list:\n"
    ),
    "plain_blockers_none": (
        "**What stands in the way.** Nothing. Every statement in every program "
        "assessed is one the converter already handles.\n"
    ),
    "plain_blocker_row": "- **{construct}** — {gloss}. Appears {count} time(s).\n",
    "plain_blocker_row_plain": "- **{construct}** — appears {count} time(s).\n",
    "plain_blockers_more": (
        "- …and {count} further construct(s), all listed in the "
        "**Unsupported-construct inventory**.\n"
    ),
    "plain_tier_note": "- **{label}** — {explanation}\n",
    "plain_effort_head": "**How much of the code this affects.**\n",
    "plain_limits": (
        "**What this report does not tell you.** It does not say what the work "
        "would cost or how long it would take, and it does not claim the "
        "converted programs would behave the same as the originals — that is "
        "measured separately, by the benchmark, not by this tool. Where a figure "
        "below is graded PLAUSIBLE rather than VERIFIED, it was derived by the "
        "documented token scan instead of a full parse of the program; "
        "appendix D explains when that happens and what it costs.\n"
    ),
}

GRADE_COLUMNS = ("Value", "Grade", "Provenance")

# --------------------------------------------------------------------------
# Plain-language layer (section 0)
# --------------------------------------------------------------------------
#
# Section 0 exists because sections 1–9 answer a reader who already knows what
# a construct is. It **translates** those findings and introduces nothing: every
# figure in it is restated from a measurement rendered below, and no sentence
# here may say anything the graded sections do not already say. That is the
# whole discipline — a plain-language summary is the easiest place in a report
# for an estimate to appear wearing the clothes of a finding (R1), so this layer
# is allowed to rephrase and forbidden to infer.

# One line per risk tier, saying what the tier means for the reader rather than
# which rule fired. Each is a plain reading of the RISK_RULES entries that can
# produce that tier — reproduced verbatim in appendix C, where the reader can
# check this wording against the policy.
TIER_IN_PLAIN_WORDS: Dict[str, Tuple[str, str]] = {
    "LOW": (
        "Ready to convert as it stands",
        "Every statement in these is one the converter already handles, and "
        "their logic is simple enough to follow end to end.",
    ),
    "MED": (
        "Convertible, after a review",
        "Something in these needs a person to look at it first — a statement "
        "the converter does not handle, a call out to another program, or "
        "logic tangled enough to be worth checking.",
    ),
    "HIGH": (
        "Substantial work before converting",
        "Either a large share of these sits outside what the converter "
        "handles, or they embed another language (such as database or "
        "transaction code) that has to be dealt with separately.",
    ),
    "BLOCKED": (
        "Cannot be converted yet",
        "Too much of the program is outside what the converter handles, or it "
        "rewrites its own control flow while running — which cannot be worked "
        "out from the source at all.",
    ),
}

# What an unsupported COBOL construct *is*, for a reader who does not write
# COBOL. Facts about the language, not claims about Relian — the counts beside
# them are the measurements. A construct with no entry is shown without a gloss
# rather than with a guessed one.
CONSTRUCT_IN_PLAIN_WORDS: Dict[str, str] = {
    "ALTER": "changes where a jump goes while the program is running",
    "CALL": "hands control to a separate program",
    "CANCEL": "unloads a separate program from memory",
    "CLOSE": "finishes with a data file",
    "DELETE": "removes a record from a data file",
    "DIVIDE": "division written as its own statement",
    "ENTRY": "declares an extra entry point into the program",
    "EXEC": "embedded database or transaction-system code",
    "GO": "jumps to another part of the program",
    "MERGE": "merges sorted data files",
    "MULTIPLY": "multiplication written as its own statement",
    "OPEN": "makes a data file ready to use",
    "READ": "reads a record from a data file",
    "RELEASE": "hands a record to a sort",
    "RETURN": "takes a record back from a sort",
    "REWRITE": "replaces a record already in a data file",
    "SORT": "sorts a data file",
    "START": "positions a data file at a particular record",
    "STRING": "joins pieces of text together",
    "SUBTRACT": "subtraction written as its own statement",
    "UNSTRING": "splits text into pieces",
    "WRITE": "adds a record to a data file",
}

# How many programs to name inline before summarising the rest, and how many
# constructs section 0 lists before deferring to the full inventory. Both
# truncations are stated in the rendered text — a silent cap reads as
# completeness.
_NAMES_SHOWN = 4
_CONSTRUCTS_SHOWN = 5

# Tiers worst-first, so the reader meets the problems before the clean results.
_TIER_ORDER = ("BLOCKED", "HIGH", "MED", "LOW")


def _name_list(names: Sequence[str], limit: int = _NAMES_SHOWN) -> str:
    """Program names for a table cell, with any truncation stated, not implied."""
    if not names:
        return "—"
    if len(names) <= limit:
        return ", ".join(f"`{n}`" for n in names)
    shown = ", ".join(f"`{n}`" for n in names[:limit])
    return f"{shown} + {len(names) - limit} more"

# Vocabulary the report must never contain (R11): commercial commitments the
# tool is not entitled to make.
FORBIDDEN_VOCABULARY = (
    "engagement fee", "fixed price", "we will deliver", "delivery date",
    "timeline", "behavioral equivalence rate", "behavioural equivalence rate",
    "guaranteed", "warranty",
)


# --------------------------------------------------------------------------
# Markdown assembly helpers
# --------------------------------------------------------------------------


def _row(*cells: object) -> str:
    return TEMPLATES["table_row"].format(cells=" | ".join(str(c) for c in cells))


def _table(headers: Sequence[str], rows: Sequence[Sequence[object]]) -> str:
    out = [_row(*headers), _row(*("---" for _ in headers))]
    out.extend(_row(*r) for r in rows)
    return "".join(out) + "\n"


def _m_row(label: str, m: Optional[Measured], absent_reason: str = "not measured") -> List[object]:
    if m is None:
        return [label, "—", "—", absent_reason]
    return [label, m.value, m.grade, m.provenance]


def _grade_line(grade: str, provenance: str) -> str:
    return TEMPLATES["grade_line"].format(grade=grade, provenance=provenance)


# --------------------------------------------------------------------------
# Migration-scope split
# --------------------------------------------------------------------------


def compute_scope_split(
    programs: Sequence[ProgramAssessment],
) -> Tuple[Optional[Measured], Optional[Measured], Tuple[Tuple[str, int], ...]]:
    """Quotable-today LOC vs LOC requiring grammar expansion, plus by-construct counts.

    Attribution is by *source line*: a code line is "requires grammar expansion"
    if it carries at least one construct C1 cannot transpile. Lines are counted
    once even when they carry several such constructs. Programs whose coverage
    could not be measured contribute to neither figure and are named in the
    provenance string — they are not silently counted as quotable.
    """
    measurable = [p for p in programs if p.coverage.coverage_ratio is not None]
    if not measurable:
        return None, None, ()

    total_code = 0
    blocked_lines = 0
    by_construct: Dict[str, int] = {}
    for p in measurable:
        total_code += p.loc.physical - p.loc.comment - p.loc.blank
        lines = {h.line for h in p.coverage.unsupported_inventory}
        blocked_lines += len(lines)
        for h in p.coverage.unsupported_inventory:
            by_construct[h.verb] = by_construct.get(h.verb, 0) + 1

    quotable = total_code - blocked_lines
    excluded = len(programs) - len(measurable)
    grades = {p.coverage.coverage_ratio.grade for p in measurable}
    grade = "VERIFIED" if grades == {"VERIFIED"} else "PLAUSIBLE"
    suffix = f"; {excluded} program(s) excluded (coverage not measured)" if excluded else ""

    quotable_m = Measured(
        quotable,
        f"code lines ({total_code}) minus lines carrying an unsupported construct "
        f"({blocked_lines}) across {len(measurable)} program(s){suffix}",
        grade,
    )
    grammar_m = Measured(
        blocked_lines,
        f"distinct code lines carrying >=1 construct outside {registry_provenance()} "
        f"across {len(measurable)} program(s){suffix}",
        grade,
    )
    ranked = tuple(sorted(by_construct.items(), key=lambda kv: (-kv[1], kv[0])))
    return quotable_m, grammar_m, ranked


def tool_versions() -> Tuple[Tuple[str, str], ...]:
    try:
        import antlr4                                   # noqa: PLC0415
        antlr_version = getattr(antlr4, "__version__", "unknown")
    except Exception:
        antlr_version = "not installed"
    try:
        import docx                                     # noqa: PLC0415
        docx_version = getattr(docx, "__version__", "installed")
    except Exception:
        docx_version = "not installed"
    return tuple(sorted({
        "python": platform.python_version(),
        "platform": platform.system(),
        "antlr4-python3-runtime": antlr_version,
        "python-docx": docx_version,
        "relian_transpiler": registry_provenance(),
        "schema": SCHEMA_VERSION,
        "cli": " ".join(Path(a).name for a in sys.argv[:1]) or "python -m src.assessment.cli",
    }.items()))


# --------------------------------------------------------------------------
# Markdown
# --------------------------------------------------------------------------


PLAIN_TITLE = "What this means"

# Labels for the two line counts section 0 restates from the migration-scope
# section. Only the wording differs; the Measured objects are the same ones.
_EFFORT_LABELS = (
    ("Lines the converter handles today", "quotable_loc"),
    ("Lines needing new converter capability first", "grammar_expansion_loc"),
)


def plain_summary(bundle: AssessmentBundle, root_label: str,
                  scope_by_construct: Sequence[Tuple[str, int]] = ()) -> Dict[str, Any]:
    """The plain-language layer as data, for every surface that renders it.

    This is the single source of section 0. :func:`render_plain_summary` formats
    it as Markdown and the API returns it as JSON for the Assess tab, so the two
    cannot drift into saying different things about the same run — the wording,
    the glosses and the figures are decided here, once.

    Nothing here is presentation. Program names are returned in full and
    unformatted, so each surface can lay them out and disclose its own
    truncation. Constructs are the exception: the list is cut to
    ``_CONSTRUCTS_SHOWN`` *with the number dropped carried alongside*, because
    both surfaces should show the same short list and both must say what they
    left out.

    Every figure is one already rendered in a graded section — restated, never
    recomputed. ``Measured`` values are passed through whole so the grade and
    provenance travel with the number (R1, R9).
    """
    t = TEMPLATES
    summary: Dict[str, Any] = {
        "title": PLAIN_TITLE,
        "intro": t["plain_intro"].strip(),
        "where_we_stand": None,
        "in_the_way": None,
        "how_much": None,
        "limits": None,
    }

    programs = bundle.programs
    if not programs:
        summary["scope"] = t["plain_scope_empty"].format(root=root_label).strip()
        return summary

    summary["scope"] = t["plain_scope"].format(
        programs=len(programs), root=root_label,
        files=len(bundle.inventory.records)).strip()

    findings = [p.risk for p in programs]
    counts = tier_counts(findings)
    by_tier = programs_by_tier(findings)
    summary["where_we_stand"] = {
        "heading": t["plain_verdict_head"].strip(),
        "grade": "PLAUSIBLE",
        "provenance": (
            "each count is the number of programs the RISK_RULES policy "
            "(appendix C) placed in that tier, restated from the Risk tiers "
            "section; the measurements the policy reads are VERIFIED"
        ),
        "groups": [
            {
                "tier": tier,
                "label": TIER_IN_PLAIN_WORDS[tier][0],
                "explanation": TIER_IN_PLAIN_WORDS[tier][1],
                "programs": counts[tier],
                "program_ids": list(by_tier.get(tier, ())),
            }
            for tier in _TIER_ORDER
            if counts.get(tier, 0)
        ],
    }

    if not scope_by_construct:
        summary["in_the_way"] = {
            "heading": t["plain_blockers_none"].strip(),
            "constructs": [],
            "omitted": 0,
        }
    else:
        summary["in_the_way"] = {
            "heading": t["plain_blockers_head"].strip(),
            "constructs": [
                {
                    "construct": construct,
                    "gloss": CONSTRUCT_IN_PLAIN_WORDS.get(construct.upper()),
                    "count": count,
                }
                for construct, count in scope_by_construct[:_CONSTRUCTS_SHOWN]
            ],
            "omitted": max(0, len(scope_by_construct) - _CONSTRUCTS_SHOWN),
        }

    summary["how_much"] = {
        "heading": t["plain_effort_head"].strip(),
        "rows": [
            {"label": label, "measured": measured_to_dict(getattr(bundle, attr))}
            for label, attr in _EFFORT_LABELS
        ],
    }
    summary["limits"] = t["plain_limits"].strip()
    return summary


def render_plain_summary(bundle: AssessmentBundle, root_label: str,
                         scope_by_construct: Sequence[Tuple[str, int]] = ()) -> str:
    """Section 0 — the findings below, in ordinary words.

    Markdown formatting of :func:`plain_summary`, which decides what it says.
    The two numeric tables carry the grade of the section they restate — the
    tier counts are PLAUSIBLE because a tier is a policy, and the line counts
    are PLAUSIBLE or VERIFIED exactly as their source measurement is.
    """
    t = TEMPLATES
    data = plain_summary(bundle, root_label, scope_by_construct)
    out: List[str] = [
        t["h1"].format(n=0, title=data["title"]),
        t["para"].format(text=data["intro"]),
        t["para"].format(text=data["scope"]),
    ]

    stands = data["where_we_stand"]
    if stands is None:                      # no programs; scope said so already
        return "".join(out)

    out.append(t["para"].format(text=stands["heading"]))
    out.append(_grade_line(stands["grade"], stands["provenance"]))
    out.append(_table(
        ("Group", "Programs", "Which ones"),
        [
            [g["label"], g["programs"], _name_list(g["program_ids"])]
            for g in stands["groups"]
        ],
    ))
    # The explanations sit under the table rather than inside it: a cell holding
    # a paragraph is a table nobody reads.
    for group in stands["groups"]:
        out.append(t["plain_tier_note"].format(
            label=group["label"], explanation=group["explanation"]))
    out.append("\n")

    blockers = data["in_the_way"]
    out.append(t["para"].format(text=blockers["heading"]))
    for item in blockers["constructs"]:
        key = "plain_blocker_row" if item["gloss"] else "plain_blocker_row_plain"
        out.append(t[key].format(
            construct=item["construct"], gloss=item["gloss"], count=item["count"]))
    if blockers["omitted"]:
        out.append(t["plain_blockers_more"].format(count=blockers["omitted"]))
    if blockers["constructs"]:
        out.append("\n")

    how_much = data["how_much"]
    out.append(t["para"].format(text=how_much["heading"]))
    out.append(_table(
        ("Measure", *GRADE_COLUMNS),
        [_m_row(row["label"], getattr(bundle, attr))
         for row, (_label, attr) in zip(how_much["rows"], _EFFORT_LABELS)],
    ))

    out.append(t["para"].format(text=data["limits"]))
    return "".join(out)


def render_markdown(bundle: AssessmentBundle, root_label: str,
                    scope_by_construct: Sequence[Tuple[str, int]] = ()) -> str:
    t = TEMPLATES
    out: List[str] = [
        t["title"].format(root=root_label),
        t["subtitle"].format(schema=bundle.schema_version,
                             manifest_hash=bundle.inventory.manifest_hash),
    ]

    cov = bundle.portfolio_coverage
    loc_totals = loc_mod.portfolio_totals([p.loc for p in bundle.programs])

    # 0. What this means — plain language, for a reader who does not write COBOL.
    out.append(render_plain_summary(bundle, root_label, scope_by_construct))

    # 1. Executive summary
    out.append(t["h1"].format(n=1, title="Executive summary"))
    out.append(_table(
        ("Measure", *GRADE_COLUMNS),
        [
            _m_row("Portfolio construct coverage", cov.coverage_ratio,
                   cov.error or "no statements recovered"),
            _m_row("Quotable-today code lines", bundle.quotable_loc),
            _m_row("Code lines requiring grammar expansion", bundle.grammar_expansion_loc),
        ],
    ))
    out.append(_grade_line(
        "PLAUSIBLE",
        "portfolio risk tier is a policy decision from the RISK_RULES table "
        "reproduced in the appendix; its inputs are VERIFIED measurements",
    ))
    out.append(_table(
        ("Measure", "Value"),
        [
            ["Portfolio risk tier", bundle.portfolio_risk.tier.value],
            ["Rule that fired", f"`{bundle.portfolio_risk.rule}`"],
        ],
    ))

    # 2. Manifest
    out.append(t["h1"].format(n=2, title="Manifest"))
    out.append(_grade_line(
        "VERIFIED",
        "sha256 and size_bytes are of the raw bytes on disk; the manifest hash is "
        "sha256 of the canonical JSON of the sorted record list "
        f"(= {bundle.inventory.manifest_hash})",
    ))
    out.append(_table(
        ("Path", "Kind", "Bytes", "Line ending", "sha256"),
        [
            [r.rel_path_posix, r.kind, r.size_bytes, r.line_ending, f"`{r.sha256[:16]}`"]
            for r in bundle.inventory.records
        ],
    ))

    # 3. LOC inventory
    out.append(t["h1"].format(n=3, title="LOC inventory"))
    out.append(_grade_line(
        "VERIFIED",
        "line categories counted per the rules in appendix A; logical statements "
        "come from the same extraction as the coverage map, and are absent where "
        "no statements could be recovered",
    ))
    out.append(_table(
        ("Program", "Physical", "Comment", "Blank", "Code", "Logical", "Method", "Dead paragraphs"),
        [
            [
                p.program_id, p.loc.physical, p.loc.comment, p.loc.blank,
                p.loc.physical - p.loc.comment - p.loc.blank,
                p.loc.logical if p.loc.logical is not None else "—",
                p.loc.logical_method,
                ", ".join(p.loc.dead_paragraphs) or "—",
            ]
            for p in bundle.programs
        ],
    ))
    out.append(t["para"].format(text=(
        f"Portfolio totals — physical {loc_totals['physical']}, "
        f"code {loc_totals['code']}, comment {loc_totals['comment']}, "
        f"blank {loc_totals['blank']}, logical "
        f"{loc_totals['logical'] if loc_totals['logical'] is not None else '—'} "
        f"({loc_totals['logical_programs_measured']} program(s) measured, "
        f"{loc_totals['logical_programs_unmeasured']} not measured)."
    )))
    notes = [p for p in bundle.programs if p.loc.note]
    if notes:
        out.append(t["h2"].format(title="LOC notes"))
        for p in notes:
            out.append(t["bullet"].format(text=f"`{p.program_id}` — {p.loc.note}"))
        out.append("\n")

    # 4. Coverage map
    out.append(t["h1"].format(n=4, title="Coverage map"))
    out.append(_table(
        ("Program", *GRADE_COLUMNS),
        [_m_row(p.program_id, p.coverage.coverage_ratio,
                p.coverage.error or "no statements recovered")
         for p in bundle.programs],
    ))
    out.append(t["h2"].format(title="Portfolio"))
    out.append(_table(
        ("Measure", *GRADE_COLUMNS),
        [
            _m_row("Coverage ratio", cov.coverage_ratio, cov.error or "not measured"),
        ],
    ))
    if cov.parser_errors:
        out.append(t["h2"].format(title="Programs excluded from the ratio"))
        for e in cov.parser_errors:
            out.append(t["bullet"].format(text=e))
        out.append("\n")

    # 5. Unsupported-construct inventory
    out.append(t["h1"].format(n=5, title="Unsupported-construct inventory"))
    ranked = cov.ranked_unsupported()
    if not ranked:
        out.append(t["para"].format(
            text="No construct outside the supported set was found."))
    else:
        out.append(_grade_line(
            "VERIFIED",
            "occurrence counts of constructs absent from "
            f"{registry_provenance()}, counted over the statements listed in the "
            "coverage map",
        ))
        out.append(_table(
            ("Construct", "Occurrences"),
            [[verb, count] for verb, count in ranked],
        ))
        out.append(t["h2"].format(title="Occurrences"))
        out.append(_table(
            ("File", "Line", "Paragraph", "Construct", "Context"),
            [
                [h.file, h.line, h.paragraph or "—", h.verb, h.context or "—"]
                for h in cov.unsupported_inventory
            ],
        ))

    # 6. DATA DIVISION features
    out.append(t["h1"].format(n=6, title="DATA DIVISION features found"))
    summary = coverage_mod.data_feature_summary([p.coverage for p in bundle.programs])
    if not summary:
        out.append(t["para"].format(text="No classified DATA DIVISION feature was found."))
    else:
        out.append(_grade_line(
            "VERIFIED",
            "occurrence counts from source; each status is probed against the "
            "transpiler itself, not asserted — `accepted_ignored` means the clause "
            "parses but is discarded, so generated code cannot depend on it",
        ))
        out.append(_table(
            ("Feature", "Occurrences", "C1 status"),
            [[f, v["occurrences"], v["status"]] for f, v in summary.items()],
        ))

    # 7. Complexity findings
    out.append(t["h1"].format(n=7, title="Complexity findings"))
    out.append(_grade_line(
        "VERIFIED",
        "computed per the formulas in appendix B; no threshold is applied here",
    ))
    out.append(_table(
        ("Program", "Cyclomatic", "Statements", "GO TO", "GO TO density",
         "ALTER", "EXEC CICS", "EXEC SQL", "Max nesting"),
        [
            [
                p.program_id,
                p.complexity.cyclomatic if p.complexity else "—",
                p.complexity.statements if p.complexity else "—",
                p.complexity.goto_count if p.complexity else "—",
                (p.complexity.goto_density.value
                 if p.complexity and p.complexity.goto_density else "—"),
                ("yes" if p.complexity.alter_present else "no") if p.complexity else "—",
                p.complexity.exec_cics_count if p.complexity else "—",
                p.complexity.exec_sql_count if p.complexity else "—",
                p.complexity.max_nesting_depth if p.complexity else "—",
            ]
            for p in bundle.programs
        ],
    ))
    fan_in = complexity_mod.copybook_fan_in(
        [p.complexity for p in bundle.programs if p.complexity]
    )
    if fan_in:
        out.append(t["h2"].format(title="Copybook fan-in"))
        out.append(_grade_line("VERIFIED", "COPY targets named in program source"))
        out.append(_table(
            ("Copybook", "Used by"),
            [[book, ", ".join(progs)] for book, progs in fan_in.items()],
        ))

    # 8. Risk tiers
    out.append(t["h1"].format(n=8, title="Risk tiers"))
    out.append(_grade_line(
        "PLAUSIBLE",
        "a published policy (RISK_RULES, appendix C), not a measurement; every "
        "input to it is VERIFIED",
    ))
    out.append(_table(
        ("Program", "Tier", "Rule that fired"),
        [[p.program_id, p.risk.tier.value, f"`{p.risk.rule}`"] for p in bundle.programs],
    ))
    out.append(_table(
        ("Tier", "Programs"),
        [[tier, str(count)] for tier, count in
         tier_counts([p.risk for p in bundle.programs]).items()],
    ))

    # 9. Migration-scope recommendation
    out.append(t["h1"].format(n=9, title="Migration-scope recommendation"))
    out.append(_table(
        ("Measure", *GRADE_COLUMNS),
        [
            _m_row("Quotable-today code lines", bundle.quotable_loc),
            _m_row("Code lines requiring grammar expansion", bundle.grammar_expansion_loc),
        ],
    ))
    out.append(t["para"].format(text=(
        "Attribution is by source line: a code line requires grammar expansion "
        "if it carries at least one construct the deterministic transpiler "
        "cannot handle. This report does not price the work and does not state "
        "a schedule."
    )))
    if scope_by_construct:
        out.append(t["h2"].format(title="By construct — what grammar work would unlock"))
        out.append(_grade_line(
            "VERIFIED",
            "occurrences of each unsupported construct across the portfolio",
        ))
        out.append(_table(
            ("Construct", "Occurrences"),
            [[verb, count] for verb, count in scope_by_construct],
        ))

    # 10. Appendices
    out.append(t["h1"].format(n=10, title="Appendices"))

    out.append(t["h2"].format(title="Appendix A — LOC counting rules"))
    out.append(t["code_block"].format(text=(loc_mod.__doc__ or "").strip()))

    out.append(t["h2"].format(title="Appendix B — complexity formulas"))
    out.append(t["code_block"].format(text=(complexity_mod.__doc__ or "").strip()))

    out.append(t["h2"].format(title="Appendix C — RISK_RULES, verbatim and in evaluation order"))
    out.append(t["code_block"].format(text="\n".join(rule_table())))

    out.append(t["h2"].format(title="Appendix D — coverage method and its limits"))
    out.append(t["code_block"].format(text=(coverage_mod.__doc__ or "").strip()))

    out.append(t["h2"].format(title="Appendix E — supported set, read from the transpiler"))
    out.append(t["para"].format(text=f"Registry: `{registry_provenance()}`"))
    out.append(t["para"].format(
        text="Supported statement keywords: " +
             ", ".join(f"`{v}`" for v in sorted(supported_verbs()))))
    out.append(t["para"].format(
        text="Statement-boundary tokens that are **not** supported: " +
             ", ".join(f"`{v}`" for v in sorted(boundary_only_tokens()))))
    out.append(_table(
        ("DATA DIVISION feature", "C1 status"),
        [[f, s] for f, s in supported_data_features().items()],
    ))

    out.append(t["h2"].format(title="Appendix F — tool versions"))
    out.append(_table(("Component", "Version"), [[k, v] for k, v in bundle.tool_versions]))

    if bundle.notes:
        out.append(t["h2"].format(title="Appendix G — notes on this run"))
        for note in bundle.notes:
            out.append(t["bullet"].format(text=note))
        out.append("\n")

    return "".join(out)


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def render_docx(markdown: str, path: Path) -> Optional[str]:
    """Write a DOCX rendering of the Markdown. Returns None (and writes nothing)
    when python-docx is unavailable, so a missing dependency never masquerades
    as a delivered artifact."""
    try:
        from docx import Document                       # noqa: PLC0415
    except ImportError:
        return None

    doc = Document()
    pending_table: List[List[str]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if not pending_table:
            return
        rows = [r for r in pending_table if not all(set(c.strip()) <= {"-"} for c in r)]
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row[:len(rows[0])]):
                    table.cell(i, j).text = cell
        pending_table = []

    in_code = False
    code_lines: List[str] = []
    for line in markdown.splitlines():
        if line.startswith("```"):
            if in_code:
                doc.add_paragraph("\n".join(code_lines), style="Intense Quote")
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            pending_table.append([c.strip() for c in line.strip("|").split("|")])
            continue
        flush_table()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    flush_table()
    doc.save(str(path))
    return str(path)


# --------------------------------------------------------------------------
# render
# --------------------------------------------------------------------------


def render(bundle: AssessmentBundle, out_dir: Path, root_label: str = "",
           scope_by_construct: Sequence[Tuple[str, int]] = (),
           json_only: bool = False, docx: bool = True) -> Dict[str, Optional[str]]:
    """Write the report artifacts. Returns paths plus the ledger hash."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    json_text = canonical_json(bundle.to_dict())
    json_bytes = json_text.encode("utf-8")
    report_hash = hashlib.sha256(json_bytes).hexdigest()

    json_path = out_dir / "assessment.json"
    json_path.write_bytes(json_bytes)
    (out_dir / "assessment.sha256").write_text(f"{report_hash}  assessment.json\n")

    result: Dict[str, Optional[str]] = {
        "json": str(json_path),
        "report_hash": report_hash,
        "md": None,
        "docx": None,
        "docx_skipped_reason": None,
    }
    if json_only:
        return result

    markdown = render_markdown(bundle, root_label or "codebase", scope_by_construct)
    md_path = out_dir / "assessment.md"
    md_path.write_text(markdown, encoding="utf-8")
    result["md"] = str(md_path)

    if not docx:
        # DOCX rendering is quadratic-ish in table rows and dominates wall time
        # on large portfolios (measured: it is the bulk of a 44-program run).
        # Skipping it is an explicit choice, recorded as one.
        result["docx_skipped_reason"] = "--no-docx was passed"
        return result

    docx_path = render_docx(markdown, out_dir / "assessment.docx")
    if docx_path is None:
        result["docx_skipped_reason"] = (
            "python-docx is not installed; no DOCX was written. "
            "Install python-docx and re-run to produce it."
        )
    result["docx"] = docx_path
    return result
